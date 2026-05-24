"""
İZÇİ — Yüksek Lisans Dönem Projesi Tez Makalesi Üretici
Bu betik, kişi istihbarat alt sistemi hakkında kapsamlı bir Word dokümanı üretir.
Çıktı: izci_tez_makalesi.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ────────────────────────────────────────────────────────────────────────
# Stil yardımcıları
# ────────────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._r.append(fld)


def style_heading(run, size, color="1f2937", bold=True):
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.font.name = "Calibri"


def add_heading(doc, text, level, color="1e3a8a"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    run = p.add_run(text)
    style_heading(run, sizes.get(level, 11), color=color)
    return p


def add_body(doc, text, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if italic:
        run.italic = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = text
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    return p


def add_code(doc, text):
    """Kod bloğu: gri arka plan, monospace yazı tipi."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F3F4F6')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("111827")
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("4b5563")


# ────────────────────────────────────────────────────────────────────────
# Doküman üretimi
# ────────────────────────────────────────────────────────────────────────
doc = Document()

# Sayfa kenar boşlukları
for sec in doc.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

# Varsayılan font
style = doc.styles['Normal']
style.font.name = "Calibri"
style.font.size = Pt(11)

# ════════════════════════════════════════════════════════════════════════
# KAPAK SAYFASI
# ════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
run = p.add_run("İZÇİ")
run.bold = True
run.font.size = Pt(48)
run.font.color.rgb = RGBColor.from_string("1e3a8a")
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("İstihbarat • Zekâ • Çözümleme • İzleme")
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor.from_string("64748b")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run(
    "Açık Kaynak Verilerden Yapay Zekâ Destekli Kişi İstihbaratı için\n"
    "Modüler ve Çoklu-Kaynak Birleştirici bir Platform"
)
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string("0f172a")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
run = p.add_run("Yüksek Lisans Dönem Projesi")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string("475569")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run("2026")
run.font.size = Pt(13)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# ÖZET
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "ÖZET", 1, color="1e3a8a")

abstract = (
    "Bu çalışma, açık kaynak haber akışlarından, sosyal medya gönderilerinden ve "
    "kamuya açık raporlardan Türkçe ve İngilizce dilinde kişi merkezli istihbarat "
    "üretmek üzere geliştirilen İZÇİ adlı yapay zekâ destekli platformu tanıtmaktadır. "
    "İZÇİ, yerel olarak çalıştırılan açık-ağırlıklı büyük dil modelleri (Gemma 3 ailesi, "
    "LLaVA 7B) aracılığıyla yapılandırılmamış metinlerden kişi, lokasyon, görüşme ve "
    "ilişki bilgisi çıkarır; çoklu-kaynak tekilleştirme algoritmasıyla aynı kişiye ait "
    "farklı bahsetmeleri eşleştirir; ve haber, görüşme, rapor ve metinde birlikte geçme "
    "kanıtlarını birleştirerek beş bağımsız sinyalden tek bir ilişki grafiği inşa eder. "
    "Üç boyutlu küre üzerinde sürekli-zamanlı bir oynatma arayüzü, kişilerin geçmişteki "
    "hareketlerini akışkan bir şekilde görüntülemeye ve trayektori-tabanlı bir tahmin "
    "modülü ile +12 / +24 / +48 / +72 saatlik gelecek konumlarını izlemeye olanak verir. "
    "Bu makale, sistemin yalnızca kişi istihbarat alt sistemini ele almakta; mimari "
    "kararları, veri modelini, modüllerin işleyişini, yapay zekâ bileşenlerini ve "
    "kullanıcı arayüzünü ayrıntılı olarak sunmaktadır."
)
p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
run = p.add_run(abstract)
run.font.size = Pt(11)
run.italic = True

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Anahtar Kelimeler: ")
run.bold = True
run.font.size = Pt(11)
run = p.add_run(
    "Açık Kaynak İstihbaratı (OSINT), Büyük Dil Modelleri, Varlık Çıkarımı, "
    "Çoklu-Kaynak Tekilleştirme, İlişki Grafiği, Konum Tahmini, Yapay Zekâ"
)
run.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# İÇİNDEKİLER
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "İÇİNDEKİLER", 1, color="1e3a8a")

toc = [
    ("1. GİRİŞ", 4),
    ("    1.1 Problemin Tanımı", 4),
    ("    1.2 Motivasyon ve Türkçe Odaklı Boşluk", 4),
    ("    1.3 Çalışmanın Katkıları", 5),
    ("    1.4 Makalenin Yapısı", 5),
    ("2. LİTERATÜR VE ARKA PLAN", 6),
    ("    2.1 Açık Kaynak İstihbaratı Pratikleri", 6),
    ("    2.2 Yapılandırılmış Çıkarım için LLM Yaklaşımları", 6),
    ("    2.3 Çoklu-Kaynak Varlık Tekilleştirme", 7),
    ("3. SİSTEM MİMARİSİ", 8),
    ("    3.1 Üç Katmanlı Genel Mimari", 8),
    ("    3.2 Teknoloji Yığını", 9),
    ("    3.3 Veri Akış Hattı", 10),
    ("4. VERİ MODELİ", 11),
    ("    4.1 Person Tablosu — Sistemin Kalbi", 11),
    ("    4.2 PersonLocation — Zaman-Konum Geçmişi", 12),
    ("    4.3 PersonMeeting — Görüşme Olayları", 12),
    ("    4.4 PersonRelationship — Yapısal İlişkiler", 13),
    ("    4.5 Destekleyici Tablolar", 13),
    ("5. MODÜLLER", 14),
    ("    5.1 Veri Toplama Modülü", 14),
    ("    5.2 Varlık Çıkarım Modülü", 15),
    ("    5.3 İşleme Hattı (Pipeline)", 16),
    ("    5.4 Tekilleştirme Modülü", 17),
    ("    5.5 Birleşik İlişki Grafiği", 18),
    ("    5.6 Zaman Çizelgesi ve Sürekli Akış", 19),
    ("    5.7 Konum Tahmini Modülü", 20),
    ("    5.8 Yüz ve Sahne Analizi", 21),
    ("    5.9 Derin Profil Sentezi", 22),
    ("6. YAPAY ZEKÂ BİLEŞENLERİ", 23),
    ("    6.1 İki Katmanlı LLM Stratejisi", 23),
    ("    6.2 JSON Şema Onarımı", 24),
    ("    6.3 Prompt Mühendisliği", 24),
    ("7. KULLANICI ARAYÜZÜ", 25),
    ("    7.1 Kişiler ve Analiz Sayfası", 25),
    ("    7.2 Birleşik Zaman Çizelgesi", 26),
    ("    7.3 Ağ Görselleştirmesi", 27),
    ("8. ETİK, GÜVENLİK VE GİZLİLİK", 28),
    ("9. SINIRLAMALAR VE GELECEK ÇALIŞMALAR", 29),
    ("10. SONUÇ", 30),
    ("KAYNAKÇA", 31),
]
for entry, _page in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(entry)
    run.font.size = Pt(11)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 1. GİRİŞ
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. GİRİŞ", 1, color="1e3a8a")

# 1.1
add_heading(doc, "1.1 Problemin Tanımı", 2)
add_body(doc,
    "Modern istihbarat çalışmasının en temel zorluklarından biri, açık kaynaklarda "
    "dağılmış halde bulunan büyük hacimli, yapılandırılmamış metinlerden belirli kişilere "
    "ait tutarlı ve birleşik bir bilgi tabanı çıkarmaktır. Bir analist, günde yüzlerce "
    "haber makalesi, sosyal medya gönderisi, basın açıklaması ve resmi raporu izlemek; "
    "bu metinlerde geçen kişileri tespit etmek; aynı kişiye ait farklı yazılışları, "
    "takma adları ve unvanları eşleştirmek; ve zamanla yayılı bir profil inşa etmek "
    "zorundadır. Bu sürecin manuel yürütülmesi yalnızca ölçeklenmemekle kalmaz, aynı "
    "zamanda tutarsız ve gözden kaçan bağlantılarla dolu bir çıktıya yol açar.")
add_body(doc,
    "Daha somut bir örnekle düşünelim: Türkiye'nin dış politikasında öne çıkan bir "
    "diplomatın son altı ayda hangi şehirlerde bulunduğunu, hangi muadilleriyle "
    "görüştüğünü ve hangi konularda demeç verdiğini öğrenmek isteyen bir analistin "
    "geleneksel yöntemlerle yapacağı iş, AA, Hürriyet, Sabah, NTV, CNN Türk gibi "
    "kaynaklarda kişinin adıyla manuel arama yapmak; her sonucu okuyup not almak; "
    "tarih ve yer bilgilerini elektronik tabloya işlemek; ve ardından bu manuel "
    "kayıtlardan zihinde bir zaman çizelgesi inşa etmektir. Bu süreç en deneyimli "
    "analistler için bile saatlerce sürer.")

# 1.2
add_heading(doc, "1.2 Motivasyon ve Türkçe Odaklı Boşluk", 2)
add_body(doc,
    "Türkçe içerik özelinde bu sorun belirgin biçimde keskinleşir. Mevcut çoğu OSINT "
    "platformu (Maltego, Palantir Gotham, Bellingcat aletleri vb.) İngilizce kaynaklı "
    "içeriği önceliklendirir ve Türkçe morfolojik karmaşıklığa karşı zayıftır. "
    "Türkçenin eklemeli yapısı (\"Erdoğan'ın\", \"Erdoğan'a\", \"Erdoğan'ı\") basit "
    "kelime eşleştirme yaklaşımlarını yetersiz kılar; takma adlar ve unvanlar "
    "(\"Cumhurbaşkanı Erdoğan\", \"RTE\", \"Recep Tayyip Erdoğan\") aynı kişiye işaret "
    "ettiği halde farklı kayıtlar olarak görünür. Türkiye odaklı analizler için "
    "Anadolu Ajansı, Hürriyet, Sabah, NTV, CNN Türk, TRT, Habertürk, Sözcü gibi yerel "
    "kaynakların entegre edilmesi ve Türkçe metinlerin doğru çıkarılması zorunludur.")
add_body(doc,
    "Buna ek olarak, son nesil büyük dil modellerinin (LLM) yerel olarak "
    "çalıştırılabilir hale gelmesi (Ollama platformu üzerinden Gemma 3, Llama 3, "
    "Mistral gibi açık-ağırlıklı modeller) bulut tabanlı API çağrılarına olan "
    "bağımlılığı ortadan kaldırmıştır. Bu, hem istihbarat verisinin üçüncü taraf "
    "sunucularına gönderilmesi gibi mahremiyet kaygılarını giderir hem de saatte "
    "yüzlerce haber işleyebilen sürekli bir hattın maliyetsiz çalışmasını sağlar. "
    "İZÇİ, bu teknoloji kavşağında doğmuş; Türkçeyi birinci sınıf vatandaş olarak "
    "ele alan, yerel LLM kullanan ve açık kaynaklı bileşenlerden kurulu bir "
    "çalışma ortamı sunmaktadır.")

# 1.3
add_heading(doc, "1.3 Çalışmanın Katkıları", 2)
add_body(doc, "Bu çalışmanın özgün katkıları aşağıdaki gibi sıralanabilir:")
contribs = [
    "Türkçe ve İngilizce çoklu kaynak akışlarını paralel olarak işleyebilen, retry "
    "ve key-rotasyon mekanizmaları ile dirençli bir veri toplama hattı.",
    "Yerel olarak çalıştırılan açık-ağırlıklı büyük dil modeli ile, yapılandırılmış "
    "JSON şemalı kişi/lokasyon/görüşme/ilişki çıkarımı.",
    "DBSCAN-benzeri kümeleme + LLM tabanlı doğrulama hibrit yöntemiyle çoklu-kaynak "
    "kişi tekilleştirme.",
    "Beş farklı sinyalden (haberde birlikte geçme, görüşme, rapor co-occurrence, "
    "ilişki tablosu, LLM derin analizi) derlenen birleşik ilişki grafiği inşası.",
    "Üç boyutlu küre üzerinde sürekli-zamanlı, requestAnimationFrame tabanlı "
    "interaktif kişi hareket akışı; gözle görülür interpolasyon, parlayan iz izi "
    "ve birim sembolojisi.",
    "Trayektori vektörü tabanlı, LLM bağımlılığı olmayan, deterministik konum "
    "tahmini modülü.",
    "Yüksek lisans seviyesinde reproducible, tek-bilgisayarda çalıştırılabilir, "
    "modüler bir referans implementasyonu.",
]
for c in contribs:
    add_bullet(doc, c)

# 1.4
add_heading(doc, "1.4 Makalenin Yapısı", 2)
add_body(doc,
    "Makalenin geri kalanı şu şekilde organize edilmiştir: Bölüm 2'de açık kaynak "
    "istihbaratı, LLM tabanlı varlık çıkarımı ve tekilleştirme literatüründeki temel "
    "yaklaşımlar incelenir. Bölüm 3'te sistemin genel mimarisi, kullanılan teknoloji "
    "yığını ve veri akışı sunulur. Bölüm 4'te ilişkisel veri modeli ve PostGIS uzamsal "
    "uzantıları açıklanır. Bölüm 5'te dokuz alt-modül tek tek ayrıntılandırılır. "
    "Bölüm 6'da yapay zekâ bileşenlerinin tasarım kararları (iki-katmanlı model, JSON "
    "şema onarımı, prompt mühendisliği) ele alınır. Bölüm 7, son kullanıcıya yönelik "
    "arayüz tasarımını gösterir. Bölüm 8 etik ve gizlilik konularını, Bölüm 9 "
    "sınırlamaları ve gelecek çalışmaları, Bölüm 10 ise genel bir sonuç sunar.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 2. LİTERATÜR VE ARKA PLAN
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. LİTERATÜR VE ARKA PLAN", 1, color="1e3a8a")

add_heading(doc, "2.1 Açık Kaynak İstihbaratı Pratikleri", 2)
add_body(doc,
    "Açık Kaynak İstihbaratı (Open Source Intelligence — OSINT), kamuya açık "
    "verilerden disipline edilmiş analiz metodları kullanarak istihbarat üretme "
    "pratiğidir. Bellingcat'in Boeing 17 olayını, Skripal zehirlenmesini ve İdlib "
    "saldırılarını dökümanlama çalışmaları, OSINT'in modern istihbaratın artık "
    "ayrılmaz bir parçası olduğunu göstermiştir. Maltego (Paterva) ve Palantir "
    "Gotham gibi ticari platformlar, görsel grafik tabanlı analiz arayüzleri "
    "sunmaktadır; ancak bu araçların lisanslama maliyetleri akademik bir bağlamda "
    "çoğu zaman erişilemezdir.")
add_body(doc,
    "Akademik literatürde Pastor-Galindo ve arkadaşları (2020), OSINT'in Kerterolojik "
    "ve metodolojik sınıflandırmasını sunmuş; Glassman ve Kang (2012) ise OSINT'in "
    "demokratik istihbarat ekosistemi içindeki rolünü tartışmıştır. Bu çalışmalar, "
    "OSINT'in dört aşamalı klasik döngüsünü tanımlar: Toplama → Filtreleme → "
    "Analiz → Yayım. İZÇİ bu döngünün ilk üç aşamasını otomasyon altına alır; "
    "yayım aşaması (insan-okunabilir raporlar) analist denetimine bırakılmıştır.")

add_heading(doc, "2.2 Yapılandırılmış Çıkarım için LLM Yaklaşımları", 2)
add_body(doc,
    "Yapay zekâ destekli varlık çıkarımının ilk nesil yaklaşımı kuralcı NER (Named "
    "Entity Recognition) sistemleridir: spaCy, Stanford NER, NLTK gibi araçlar, "
    "önceden eğitilmiş istatistiksel modellerle metinden kişi, kurum, lokasyon "
    "varlıklarını ayıklar. Bu yaklaşım hızlıdır, ancak iki kritik sınırlaması "
    "vardır: (i) Türkçe gibi düşük-kaynaklı diller için modeller İngilizceye göre "
    "belirgin biçimde zayıftır; (ii) sadece varlığı tespit eder, varlıklar "
    "arasındaki ilişkileri (\"X ile Y görüştü\", \"Z'nin Ankara'daki ofisinde "
    "bulundu\") çıkarmaz.")
add_body(doc,
    "Modern büyük dil modelleri (BERT, GPT-3, Llama, Gemma ailesi) bu sınırlamaları "
    "önemli ölçüde aşar. JSON şema kısıtlı çıktı üretme yeteneği sayesinde, tek bir "
    "prompt aracılığıyla \"kim, kim ile, nerede, ne zaman, ne yaptı\" sorularına "
    "yapılandırılmış yanıt alınabilir. İZÇİ bu paradigmayı benimser; yerel olarak "
    "Ollama platformu üzerinde barındırılan Gemma 3:12B (büyük çıkarımlar) ve Gemma "
    "3:4B (hızlı doğrulamalar) modellerini kullanır. Bu seçim, hem mahremiyeti hem "
    "de maliyet bağımsızlığını korur.")

add_heading(doc, "2.3 Çoklu-Kaynak Varlık Tekilleştirme", 2)
add_body(doc,
    "Aynı gerçek-dünya kişisinin farklı haberlerde farklı yazılışlarla geçmesi, "
    "naif kelime eşleştirmesinin başarısızlığına yol açar. Klasik yaklaşımlar "
    "Levenshtein mesafesi, Jaro-Winkler benzerliği veya phonetic encoding "
    "(Soundex, Metaphone) gibi karakter-tabanlı metrikler kullanır. Bu yöntemler "
    "\"Erdoğan\" ile \"Erdogan\" gibi imla varyasyonlarını yakalayabilir, ancak "
    "\"Cumhurbaşkanı\", \"RTE\", \"Reis\" gibi tamamen farklı kelime kümeleriyle "
    "ifade edilen aynı kişiyi yakalayamaz.")
add_body(doc,
    "İZÇİ, bu sorunu iki aşamalı bir hibrit yaklaşımla çözer: birinci aşamada "
    "yüksek güvenli karakter benzerliği aday çiftleri üretilir; ikinci aşamada "
    "belirsiz adaylar küçük bir LLM'e (\"Bu iki isim aynı kişiye mi ait?\") "
    "doğrulatılır. Bu hibrit yaklaşım, deterministik hız ve LLM bilgi-zenginliği "
    "arasında denge kurar.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 3. SİSTEM MİMARİSİ
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. SİSTEM MİMARİSİ", 1, color="1e3a8a")

add_heading(doc, "3.1 Üç Katmanlı Genel Mimari", 2)
add_body(doc,
    "İZÇİ üç ayrı katman üzerinde işler. Her katmanın sorumlulukları ve diğer "
    "katmanlarla arayüzleri açıkça tanımlanmıştır; bu, modüllerin bağımsız olarak "
    "geliştirilmesini ve test edilmesini mümkün kılar.")

# Mimari tablo
add_body(doc, "Şekil 1, üç katmanın özet sorumluluklarını göstermektedir.")
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
header = table.rows[0].cells
header[0].text = "Katman"
header[1].text = "Sorumluluk"
header[2].text = "Temel Bileşenler"
for c in header:
    set_cell_shading(c, "1e3a8a")
    for p in c.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor.from_string("FFFFFF")
            r.bold = True

rows = [
    ("Toplama", "Periyodik olarak çalışan veri toplayıcılar; "
                "ham haber/mesajları DB'ye yazar.",
                "APScheduler, RSS Toplayıcı, NewsAPI, Telegram"),
    ("İşleme", "LLM tabanlı çıkarım, tekilleştirme, ilişki grafiği inşası.",
                "Gemma3:12b, Gemma3:4b, Pipeline, Dedup"),
    ("Sunum",  "Analist arayüzü; REST API + 3D küre.",
                "FastAPI, Vue.js 3, Cesium.js, vis-network"),
]
for i, (a, b, c) in enumerate(rows, start=1):
    row = table.rows[i].cells
    row[0].text = a; row[1].text = b; row[2].text = c
add_caption(doc, "Şekil 1. İZÇİ üç katmanlı mimari özetinin tablo gösterimi.")

add_heading(doc, "3.2 Teknoloji Yığını", 2)
add_body(doc,
    "Sistem aşağıdaki açık kaynaklı teknolojiler üzerinde inşa edilmiştir. "
    "Bağımlılıkların seçimi sırasında üç ilkeyi gözettik: (i) yerel olarak "
    "çalıştırılabilir olmak, (ii) topluluk tarafından aktif olarak geliştirilmek, "
    "(iii) Türkçe içerikle uyumlu olmak.")

stack = [
    ("Backend Çatı", "FastAPI (asyncio destekli HTTP API), Pydantic (veri doğrulama), Uvicorn (ASGI sunucu)."),
    ("Veri Tabanı", "PostgreSQL 15 + PostGIS (jeo-uzamsal sorgular için); SQLite + SpatiaLite "
                    "geliştirme-modu yedek seçenek olarak; SQLAlchemy 2.x ORM; Alembic migrasyon."),
    ("Worker", "APScheduler BlockingScheduler — RSS, NewsAPI, Telegram, profil güncelleme, "
              "yüz analizi, tekilleştirme görevleri için zamanlanmış cronlar."),
    ("LLM", "Ollama (yerel servis) → Gemma 3:12B (\"large\"), Gemma 3:4B (\"small\"), "
            "LLaVA 7B (görsel)."),
    ("Frontend", "Vue.js 3 (single-file components yerine doğrudan template composition), "
                "Leaflet (2D harita), Cesium.js (3D küre), vis.js Network (graf), "
                "Chart.js (analitik)."),
]
for k, v in stack:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("• " + k + ": "); run.bold = True; run.font.size = Pt(11)
    run = p.add_run(v); run.font.size = Pt(11)

add_heading(doc, "3.3 Veri Akış Hattı", 2)
add_body(doc,
    "Bir haberin sisteme girişinden, kişi profilinde görünür hale gelmesine kadar "
    "olan akış aşağıda açıklanmıştır. Akış asenkron olarak çalışır; her aşama "
    "kendi tetiklenme zamanını APScheduler tarafından alır.")
flow = [
    "ADIM 1 — Toplama: RSS toplayıcı her 10 dakikada, NewsAPI 30 dakikada, Telegram "
    "15 dakikada bir tetiklenir. Her kaynak başına bağımsız bir thread'de çalışılır "
    "ve sonuçlar `News` tablosuna `is_processed=False` olarak yazılır.",
    "ADIM 2 — Çıkarım: Pipeline işi, henüz işlenmemiş haberleri ve retry kuyruğuna "
    "göre uygun olanları çeker; Gemma3:12B'ye verir; JSON yanıtı `ExtractedData` "
    "veri sınıfına dönüştürür.",
    "ADIM 3 — Yazma: ExtractedData içinden çıkan kişiler `_upsert_person` ile "
    "Person tablosunda upsert edilir; her bahsetme için PersonLocation kaydı "
    "düşülür; varsa görüşmeler PersonMeeting'e eklenir.",
    "ADIM 4 — Tetikleme: Yeni bir bahsetme alan kişinin `needs_profile_update` "
    "alanı True yapılır.",
    "ADIM 5 — Tekilleştirme: Gece çalışan dedup işi, Levenshtein + LLM hibrit "
    "kümeleme ile aynı kişiye ait farklı kayıtları birleştirir.",
    "ADIM 6 — Derin Profil: needs_profile_update=True olan kişiler için Gemma3:12B'ye "
    "son 30 günlük tüm bahsetmeler verilir; intelligence_summary, vulnerability_notes, "
    "key_relationships_json gibi alanlar üretilir.",
    "ADIM 7 — Sunum: Frontend, REST API üzerinden bu zenginleştirilmiş profili çeker "
    "ve 3D küre + 2D detay panellerinde görselleştirir.",
]
for s in flow:
    add_bullet(doc, s)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 4. VERİ MODELİ
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. VERİ MODELİ", 1, color="1e3a8a")

add_body(doc,
    "İZÇİ'nin veri modeli, kişiyi merkeze alan ve her ek bilgiyi (lokasyon, "
    "görüşme, ilişki, alias, derin analiz) buna bağlayan bir yıldız topolojisidir. "
    "Aşağıda her tablonun amacı ve önemli alanları açıklanmaktadır.")

add_heading(doc, "4.1 Person Tablosu — Sistemin Kalbi", 2)
add_body(doc,
    "Person tablosu sistemin en kritik tablosudur. Her satır bir gerçek-dünya "
    "kişisine karşılık gelir. Tekilleştirme garantisi bu tablo seviyesinde verilir. "
    "Tablo iki kavramsal grup alan içerir: (a) doğrudan toplanan ham bilgiler "
    "(name, role, last_known_city), (b) LLM tarafından zenginleştirilen analitik "
    "alanlar (intelligence_summary, predicted_movements).")

add_code(doc,
    "class Person(Base):\n"
    "    __tablename__ = 'persons'\n"
    "    id                   = Column(Integer, primary_key=True)\n"
    "    name                 = Column(String, nullable=False)\n"
    "    name_normalized      = Column(String, index=True)\n"
    "    role                 = Column(String)\n"
    "    nationality          = Column(String)\n"
    "    profile_notes        = Column(Text)\n"
    "    mention_count        = Column(Integer, default=0)\n"
    "    first_seen           = Column(DateTime)\n"
    "    last_seen            = Column(DateTime)\n"
    "    last_known_country   = Column(String)\n"
    "    last_known_city      = Column(String)\n"
    "    geom                 = Column(Geometry('POINT', 4326))   # PostGIS\n"
    "    is_turkey_related    = Column(Boolean, default=False)\n"
    "    organization_affiliation = Column(String)\n"
    "    intelligence_summary = Column(Text)        # LLM derin analiz\n"
    "    vulnerability_notes  = Column(Text)\n"
    "    strength_notes       = Column(Text)\n"
    "    predicted_movements  = Column(Text)\n"
    "    pattern_of_life_summary = Column(Text)\n"
    "    needs_profile_update = Column(Boolean, default=False)\n"
    "    key_relationships_json = Column(JSON)")
add_caption(doc, "Kod 1. Person tablosu için SQLAlchemy modeli (kısaltılmış).")

add_heading(doc, "4.2 PersonLocation — Zaman-Konum Geçmişi", 2)
add_body(doc,
    "Bir kişinin haberlerde geçtiği her bağlamda bir lokasyon çıkarılabilir. Bu "
    "lokasyon kişiye doğrudan değil, PersonLocation tablosuna ayrı bir satır olarak "
    "yazılır. Bu yapı, aynı kişinin zamanla farklı şehirlerde görünmesini takip "
    "etmek için zorunludur.")
add_body(doc,
    "Özellikle dikkat edilen iki alan: timestamp (haberin yayım zamanına bağlı) ve "
    "location_type (mention | meeting | residence | travel — kategorize edilmiş "
    "bağlam). Hareket akışı modülü, bu tablodaki sıralı kayıtlardan zaman-mekân "
    "interpolasyonu yapar.")

add_heading(doc, "4.3 PersonMeeting — Görüşme Olayları", 2)
add_body(doc,
    "İki kişi arasındaki dokumantasyona dayalı görüşmeler PersonMeeting tablosunda "
    "saklanır. person_a_id ve person_b_id alanları, simetrik ilişkinin iki "
    "tarafını referanslar. Tablo aynı zamanda meeting_type "
    "(diplomatik | iş | aile | resmi) ve context (görüşmenin doğal-dil özetı) "
    "alanlarını içerir; bu alanlar daha sonra ilişki grafiğinde sentiment çıkarımı "
    "için kullanılır.")

add_heading(doc, "4.4 PersonRelationship — Yapısal İlişkiler", 2)
add_body(doc,
    "Görüşme bireysel bir olaydır; ilişki ise süreç içinde kararlı hale gelmiş bir "
    "bağdır. PersonRelationship tablosu bu süreçsel bağı modellenir. "
    "evidence_count alanı, ilişkinin gözlemlendiği farklı kanıt sayısıdır; "
    "strength alanı bu kanıt sayısının normalize edilmiş skorudur. sentiment "
    "alanı (positive | neutral | negative) görüşme bağlamlarındaki anahtar "
    "kelimelerden çıkarılır.")

add_heading(doc, "4.5 Destekleyici Tablolar", 2)
add_body(doc,
    "Yukarıda anlatılan dört çekirdek tabloya ek olarak sistem aşağıdaki destekleyici "
    "tabloları kullanır:")
support_tables = [
    "**News** ve **news_person** join tablosu — Haber metni ve hangi kişilerin "
    "geçtiği. İlişki grafiğindeki en önemli sinyal kaynağı.",
    "**Quote** — Kişiye doğrudan atfedilen alıntılar; söylem analizi için.",
    "**ReportExtraction** — Yapılandırılmış raporlardan çıkarılan kişi-olay "
    "bağları; askeri raporlardan ve resmi bildirilerden.",
    "**FaceRecord** — LLaVA görsel analizinden gelen yüz/sahne kayıtları.",
    "**PersonAlias** — Aynı kişinin tespit edilmiş alternatif yazılışları.",
    "**PatternOfLife** ve **PersonProfileAnalysis** — Derin LLM analizinin "
    "iki ayrı çıktı tablosu.",
]
for t in support_tables:
    add_bullet(doc, t.replace("**", ""))

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 5. MODÜLLER
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. MODÜLLER", 1, color="1e3a8a")

# 5.1
add_heading(doc, "5.1 Veri Toplama Modülü", 2)
add_body(doc,
    "Veri toplama modülü `core/collectors/` dizini altında üç bağımsız toplayıcıdan "
    "oluşur. Her toplayıcı kendi konfigürasyon kaynağına ve kendi zamanlama "
    "döngüsüne sahiptir.")

add_heading(doc, "5.1.1 RSS Toplayıcı", 3)
add_body(doc,
    "RSS toplayıcı, 230+ feed (15 kategoride gruplandırılmış) ve 13 ayrıca "
    "bayraklanmış Türkiye odaklı feed'i tarar. Konfigürasyon `config.GLOBAL_RSS_FEEDS` "
    "ve `config.TURKEY_RSS_FEEDS` sözlüklerinde merkezi olarak yönetilir. Toplayıcı "
    "bir Python `ThreadPoolExecutor` (varsayılan max_workers=6) kullanır; her "
    "(feed × keyword) çifti bağımsız bir thread'e atanır ve her thread kendi DB "
    "session'ını alır — bu thread güvenliği için kritiktir çünkü SQLAlchemy "
    "session'ları thread-local değildir.")

add_heading(doc, "5.1.2 NewsAPI Toplayıcı", 3)
add_body(doc,
    "NewsAPI ücretsiz katmanı 100 istek/gün ile sınırlıdır. İZÇİ bunu aşmak için "
    "çoklu API key rotasyonu kullanır: bir key kotaya çarptığında sıradaki key'e "
    "geçer ve günlük tükenmiş key listesi tutar. Bu mekanizma, sınırlı bir bütçeyle "
    "günde binlerce haber toplamayı mümkün kılar.")

add_heading(doc, "5.1.3 Telegram Toplayıcı", 3)
add_body(doc,
    "Telegram resmi API'sinin (MTProto) ağır kimlik doğrulama gereksinimlerini "
    "atlamak için modül `t.me/s/{kanal}` HTML scrape yaklaşımını kullanır — bu, "
    "kamuya açık kanalların web önizlemesidir. BeautifulSoup ile HTML ayrıştırılır, "
    "her mesaj için topic sınıflandırma (conflict, breaking, alerts, osint, "
    "politics, middleeast) yapılır. Bu sınıflandırma erken sinyal radarı olarak "
    "frontend'in 'erken sinyal' rozetinde kullanılır.")

# 5.2
add_heading(doc, "5.2 Varlık Çıkarım Modülü", 2)
add_body(doc,
    "Toplanmış ham haberi kişi, lokasyon, görüşme, ilişki gibi yapılandırılmış "
    "varlıklara dönüştüren modül `core/llm/extractor.py` dosyasındadır. Modül, "
    "Gemma3:12B'ye bir EXTRACTION_PROMPT şablonuyla istekte bulunur ve yanıtın "
    "katı bir JSON şemasına uymasını şart koşar.")

add_code(doc,
    "{\n"
    "  \"topic_location\":      {\"country\": \"Türkiye\", \"city\": \"Ankara\"},\n"
    "  \"statement_location\":  {\"country\": \"...\",      \"city\": \"...\"},\n"
    "  \"persons\": [\n"
    "    {\n"
    "      \"name\": \"...\", \"role\": \"...\", \"nationality\": \"...\",\n"
    "      \"organization\": \"...\"\n"
    "    }\n"
    "  ],\n"
    "  \"organizations\": [\n"
    "    {\"name\": \"...\", \"type\": \"...\"}\n"
    "  ],\n"
    "  \"meetings\": [\n"
    "    {\n"
    "      \"person_a\": \"...\", \"person_b\": \"...\",\n"
    "      \"type\": \"diplomatic|business|family|official\",\n"
    "      \"location\": \"...\", \"context\": \"...\",\n"
    "      \"date\": \"YYYY-MM-DD\"\n"
    "    }\n"
    "  ],\n"
    "  \"is_turkey_relevant\": true,\n"
    "  \"relevance_score\": 0.87\n"
    "}")
add_caption(doc, "Kod 2. EXTRACTION_PROMPT'tan beklenen JSON çıktı şeması.")

add_body(doc,
    "Çıktı, Python dataclass'larına dönüştürülür: ExtractedData, ExtractedPerson, "
    "ExtractedMeeting, ExtractedLocation. Bu sayede pipeline'ın geri kalanı LLM "
    "yanıtının yapısal güvencesini varsayar.")

# 5.3
add_heading(doc, "5.3 İşleme Hattı (Pipeline)", 2)
add_body(doc,
    "İşleme hattı `core/processing/pipeline.py` içindeki `ProcessingPipeline` "
    "sınıfında tanımlıdır. Sınıf, henüz işlenmemiş News kayıtlarını batch'ler "
    "halinde işler. Pipeline'ın retry mantığı dikkatle tasarlanmıştır: bir LLM "
    "çağrısı başarısız olduğunda kayıt sonsuza kadar tekrarlanmaz; bunun yerine "
    "exponential backoff uygulanır.")

retry_table = doc.add_table(rows=4, cols=3)
retry_table.style = 'Light Grid Accent 1'
retry_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = retry_table.rows[0].cells
hdr[0].text = "Hata"; hdr[1].text = "Bekleme"; hdr[2].text = "Davranış"
for c in hdr:
    set_cell_shading(c, "1e3a8a")
    for p in c.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor.from_string("FFFFFF"); r.bold = True
data = [
    ("1. hata", "5 dakika", "Yeniden kuyruğa alınır."),
    ("2. hata", "30 dakika", "Yeniden kuyruğa alınır, attempt+=1."),
    ("3. hata", "120 dakika", "Son şans; PROCESS_MAX_ATTEMPTS=3 sonrasında kalıcı atlama."),
]
for i, (a, b, c) in enumerate(data, 1):
    row = retry_table.rows[i].cells
    row[0].text = a; row[1].text = b; row[2].text = c
add_caption(doc, "Tablo 1. Pipeline retry/backoff politikası.")

add_body(doc,
    "Bu retry stratejisi, Ollama servisi geçici olarak çöktüğünde bile haberlerin "
    "kaybolmamasını ve sistemin kendi başına toparlanmasını sağlar.")

# 5.4
add_heading(doc, "5.4 Tekilleştirme Modülü", 2)
add_body(doc,
    "Tekilleştirme `core/processing/dedup.py` içindeki iki aşamalı bir hatta "
    "yapılır. Birinci aşama hızlı bir kümelemedir; ikinci aşama belirsiz adayları "
    "LLM'e doğrulatır.")

add_heading(doc, "5.4.1 Aşama 1 — Aday Kümeleri (DBSCAN-Benzeri)", 3)
add_body(doc,
    "Her Person kaydı için bir bag-of-words vektör çıkarılır: ad, takma adlar, "
    "son lokasyonlar, organizasyon. İki kayıt arasındaki Levenshtein normalize "
    "benzerliği 0.7'nin üzerindeyse aday çift olarak işaretlenir. ε=0.7, deneysel "
    "olarak yanlış-pozitif oranını minimize edip yanlış-negatif oranını kabul "
    "edilebilir tutan değerdir.")

add_heading(doc, "5.4.2 Aşama 2 — LLM Doğrulama", 3)
add_body(doc,
    "Aşama 1'den çıkan adaylardan benzerliği 0.85'in altında kalanlar — yani "
    "\"belirsiz\" bölgede olanlar — Gemma3:4B'ye iletilir. Bu küçük model, "
    "yanıtı katı bir formatta vermesi için promptlanır:")

add_code(doc,
    "Sen iki kişi adını verilecek. Aynı kişi mi olduğunu söyle.\n"
    "İsim 1: \"{name_a}\" (rol: {role_a}, lokasyon: {city_a})\n"
    "İsim 2: \"{name_b}\" (rol: {role_b}, lokasyon: {city_b})\n"
    "\n"
    "Sadece YES veya NO ile cevap ver.")
add_caption(doc, "Kod 3. LLM doğrulama promptu.")

add_heading(doc, "5.4.3 Birleştirme", 3)
add_body(doc,
    "Doğrulanan çiftler için birleştirme yapılır: daha eski (id'si küçük olan) "
    "kayıt korunur ve yeni kaydın tüm referansları (PersonLocation, PersonMeeting, "
    "news_person, ReportExtraction) eski kişiye yönlendirilir. Bu yaklaşım eski "
    "kayıtların \"kanonik\" olmasını sağlar ve audit trail açısından "
    "tutarlılığı garanti eder.")

# 5.5
add_heading(doc, "5.5 Birleşik İlişki Grafiği", 2)
add_body(doc,
    "Bu sistemin en özgün modülü, bir kişi için ilişki grafiği inşa ederken beş "
    "ayrı sinyal kaynağını birleştiren `build_person_relationships(db, person_id)` "
    "fonksiyonudur. Her bir kaynak tek başına eksik olabilir; ancak birleşim, "
    "küçük veri kümelerinde bile anlamlı bağlantı tespiti sağlar.")

sources = [
    "**(1) PersonMeeting tablosu** — Yapılandırılmış görüşme kayıtları. En güçlü "
    "ve doğrudan sinyal; ancak tablo yalnızca açık görüşme bahsedilmesi olduğunda "
    "doldurulur.",
    "**(2) ReportExtraction co-occurrence** — Aynı resmi raporda birlikte "
    "geçen kişiler. Yapılandırılmış raporlar nadir olduğundan bu sinyal sınırlı.",
    "**(3) news_person co-mention** — En yaygın sinyal kaynağı. İki kişi aynı "
    "haber makalesinde geçiyorsa muhtemelen ilişkilidirler.",
    "**(4) PersonRelationship tablosu** — Daha önce işaretlenmiş ilişkiler "
    "(örneğin LLM derin analizinden gelen).",
    "**(5) key_relationships_json** — LLM derin analizinin doğrudan ilişki "
    "tahminleri (örneğin kişinin profil notunda \"X'in en yakın danışmanı\" "
    "yazıyorsa).",
]
for s in sources:
    add_bullet(doc, s.replace("**", ""))

add_body(doc,
    "Beş kaynak tek bir defaultdict yapısında toplanır. Her kişi-çifti için kanıt "
    "sayısı (meetings + distinct(co_reports) + distinct(co_news)) hesaplanır; "
    "strength = ev_count / max_evidence_in_dataset olacak şekilde normalize "
    "edilir. Sentiment, görüşme tipi ve bağlam metni içindeki POSITIVE_KEYWORDS "
    "(\"görüştü\", \"anlaştı\", \"destek verdi\") ve NEGATIVE_KEYWORDS "
    "(\"eleştirdi\", \"reddetti\", \"tehdit etti\") sayımlarına göre belirlenir.")

# 5.6
add_heading(doc, "5.6 Zaman Çizelgesi ve Sürekli Akış", 2)
add_body(doc,
    "Zaman çizelgesi modülü iki bileşenden oluşur: backend'de "
    "`build_deduplicated_timeline(db, person_id)` ve frontend'de sürekli zamanlı "
    "oynatma motoru.")

add_heading(doc, "5.6.1 Tekilleştirilmiş Zaman Çizelgesi", 3)
add_body(doc,
    "Aynı kişinin aynı şehirde aynı gün gerçekleşen birden fazla bahsetmesi tek bir "
    "olay satırına birleştirilir. Bu birleştirme aşağıdaki yan ürünleri üretir:")
ts_outputs = [
    "source_count — kaç farklı haberden teyit edildi",
    "all_descriptions — her kaynaktaki orjinal cümle özeti",
    "combined_confidence — kaynak sayısına göre artan güven skoru",
    "meetings — bu olayda paralel olarak kaydedilen görüşmeler",
]
for s in ts_outputs:
    add_bullet(doc, s)

add_heading(doc, "5.6.2 Sürekli Zaman Akışı (Frontend)", 3)
add_body(doc,
    "Geleneksel zaman çizelgeleri olayları ayrık simgeler olarak gösterir; İZÇİ "
    "ise sürekli zaman ekseni üzerinde 60fps akıcı oynatma sunar. Bu, "
    "`requestAnimationFrame` tabanlı bir döngüyle gerçekleştirilir.")
add_code(doc,
    "function _startFlowAnimation() {\n"
    "    const tick = (nowMs) => {\n"
    "        const dtSec = (nowMs - state.milFlow.lastFrameMs) / 1000;\n"
    "        state.milFlow.lastFrameMs = nowMs;\n"
    "        const scenarioDelta = dtSec * speed * wallTimeRatio;\n"
    "        state.milFlow.currentMs += scenarioDelta;\n"
    "        _milRenderFlowAt(state.milFlow.currentMs);\n"
    "        state.milFlow.rafId = requestAnimationFrame(tick);\n"
    "    };\n"
    "    state.milFlow.rafId = requestAnimationFrame(tick);\n"
    "}")
add_caption(doc, "Kod 4. Sürekli-zaman oynatma döngüsü (kısaltılmış).")

add_body(doc,
    "Her tick'te, her track için bracket arama + lineer interpolasyon "
    "(`_milInterpAt`) çalıştırılır. İki bitişik PersonLocation noktası "
    "arasında doğrusal olarak konum belirlenir; bu, kullanıcıya kişinin iki "
    "şehir arasında \"akışkan\" şekilde hareket ettiği yanılsamasını verir. "
    "Aynı zamanda her track için son ~4 saatlik 'iz' parlayan bir polyline "
    "olarak çizilir (`PolylineGlowMaterialProperty`).")

# 5.7
add_heading(doc, "5.7 Konum Tahmini Modülü", 2)
add_body(doc,
    "Konum tahmini, deterministik bir trayektori vektörü hesaplaması üzerine "
    "kuruludur. LLM kullanmaz; bu, hızlı ve tekrarlanabilir sonuç üretir.")

add_heading(doc, "5.7.1 Trayektori Vektörü", 3)
add_body(doc,
    "Bir kişi/birim için son iki PersonLocation kaydı arasındaki açısal hız "
    "vektörü hesaplanır:")
add_code(doc,
    "v_lat = (lat_n - lat_{n-1}) / Δt_h\n"
    "v_lng = (lng_n - lng_{n-1}) / Δt_h\n"
    "\n"
    "# Δt_h = saat cinsinden zaman farkı")
add_caption(doc, "Kod 5. Trayektori vektörü hesaplama formülü.")

add_heading(doc, "5.7.2 Geleceğe Projeksiyon", 3)
add_body(doc,
    "+12, +24, +48, +72 saat ufuklarında her birim için tahmini konum:")
add_code(doc,
    "predicted_lat = current_lat + v_lat × Δh\n"
    "predicted_lng = current_lng + v_lng × Δh\n"
    "confidence    = base × (1 - Δh / 96)")
add_caption(doc, "Kod 6. Doğrusal projeksiyon ve doğrusal güven azalımı.")

add_body(doc,
    "Güven skoru ufuk uzadıkça doğrusal olarak azalır; +72 saatte 25%'lik bir "
    "sıkı taban'a oturtulur. Eğer son iki konum yakın zamanda alınmışsa "
    "(\"trajectory\" yöntemi) base=0.85; bilinen bir yön/hız varsa ama "
    "trajektori yoksa base=0.55; tamamen sabit duruyorsa base=0.35.")

add_heading(doc, "5.7.3 Engagement Zone Tespiti", 3)
add_body(doc,
    "Birden fazla birimin tahmini konumları aynı zaman ufkunda 8 km içinde "
    "kesişiyorsa, bu bir \"olası temas bölgesi\" olarak işaretlenir. Bu mekanizma "
    "sadece askeri bağlamda anlamlıdır; intel kişileri için tetiklenmez. Çünkü "
    "iki diplomatın aynı şehirde olması rutindir; iki düşman birliğin aynı "
    "vadide olması ise kayda değer bir olaydır.")

# 5.8
add_heading(doc, "5.8 Yüz ve Sahne Analizi", 2)
add_body(doc,
    "Çoğu haber bir görsel içerir ve bu görseldeki kişiler/sahne, sadece metinden "
    "elde edilemeyen ek istihbarat sağlar. İZÇİ bunu LLaVA 7B çok-modlu modeli ile "
    "gerçekleştirir.")

add_body(doc, "LLaVA için tasarlanan prompt aşağıdaki yapılandırılmış JSON çıktıyı "
    "talep eder:")
add_code(doc,
    "{\n"
    "  \"persons_detected\":   [...],\n"
    "  \"scene_type\":         \"meeting | press | military | civilian | ...\",\n"
    "  \"war_detected\":       true | false,\n"
    "  \"scene_location_country\": \"...\",\n"
    "  \"scene_location_city\":    \"...\"\n"
    "}")
add_caption(doc, "Kod 7. LLaVA görsel analiz çıktısı şeması.")

add_body(doc,
    "Çıkan persons_detected listesi DB'deki Person kayıtlarıyla isim benzerliği "
    "üzerinden eşleştirilmeye çalışılır; başarılı eşleştirme FaceRecord'da kişi "
    "ile bağlantı kurar. war_detected=True olan görseller harita üzerinde 💥 "
    "işaretiyle gösterilir.")

# 5.9
add_heading(doc, "5.9 Derin Profil Sentezi", 2)
add_body(doc,
    "Sistemin en LLM-yoğun adımı budur. Gece 03:00'te çalışan bu cron işi, "
    "needs_profile_update=True olan tüm kişileri Gemma3:12B'ye verir; modelden "
    "kişiyi 360-derece anlamasını ve aşağıdaki alanları üretmesini ister:")
deep_outputs = [
    "intelligence_summary — kişinin son 30 günlük aktivitesinin 3-4 paragraflık "
    "yapılandırılmış özeti",
    "vulnerability_notes — zayıf noktalar (politik, ekonomik, kişisel)",
    "strength_notes — güç noktaları (etki ağı, kaynaklar)",
    "predicted_movements — yakın gelecekte muhtemel hareketler ve gerekçeler",
    "pattern_of_life_summary — günlük rutin, tercih edilen mekânlar, iletişim "
    "kalıpları",
    "key_relationships — en kritik 5-8 ilişki (isim, ilişki türü, önem)",
]
for s in deep_outputs:
    add_bullet(doc, s)

add_body(doc,
    "Bu işlemi gece çalıştırmamızın iki sebebi vardır: (i) Gemma3:12B her kişi "
    "için ortalama 8-15 saniye sürer; gündüz pipeline'ı yavaşlatır. (ii) Analist "
    "günlük raporlarını sabah açtığında zaten dünkü tüm bahsetmelerin sentezini "
    "görür.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 6. YAPAY ZEKÂ BİLEŞENLERİ
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. YAPAY ZEKÂ BİLEŞENLERİ", 1, color="1e3a8a")

add_heading(doc, "6.1 İki Katmanlı LLM Stratejisi", 2)
add_body(doc,
    "İZÇİ üç model boyutu kullanır; her model boyutu farklı bir görev sınıfına "
    "tahsis edilmiştir. Bu, hem hız hem de kalite açısından dengeli bir karışım "
    "sağlar.")

llm_table = doc.add_table(rows=4, cols=3)
llm_table.style = 'Light Grid Accent 1'
llm_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = llm_table.rows[0].cells
hdr[0].text = "Model"; hdr[1].text = "Tipik Görev"; hdr[2].text = "Tipik Yanıt Süresi"
for c in hdr:
    set_cell_shading(c, "1e3a8a")
    for p in c.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor.from_string("FFFFFF"); r.bold = True
llm_rows = [
    ("Gemma 3:12B", "Karmaşık çıkarım: haber işleme, derin profil sentezi", "8–15 sn"),
    ("Gemma 3:4B", "Hızlı doğrulama: dedup karar, basit sınıflandırma", "1–3 sn"),
    ("LLaVA 7B", "Görsel: yüz / sahne analizi", "5–10 sn"),
]
for i, (a, b, c) in enumerate(llm_rows, 1):
    r = llm_table.rows[i].cells
    r[0].text = a; r[1].text = b; r[2].text = c
add_caption(doc, "Tablo 2. LLM model dağılımı ve tipik görevler.")

add_heading(doc, "6.2 JSON Şema Onarımı", 2)
add_body(doc,
    "LLM'lerin yapılandırılmış JSON çıktısı vermesi her zaman temiz değildir: "
    "model bazen markdown code fence ile sarar, bazen kısa cevap modunda "
    "kapatmadığı parantezler bırakır. `OllamaClient.generate_json` üç aşamalı "
    "dayanıklı bir parser kullanır:")
json_steps = [
    "Aşama 1: doğrudan json.loads dene; başarısızsa devam et.",
    "Aşama 2: ```json ... ``` markdown code fence'i çıkar, tekrar dene.",
    "Aşama 3: metin içindeki ilk `{` ile son `}` arasındaki substring'i ayır; "
    "kapanmayan {[\" karakterlerini saymaya dayalı bir onarım fonksiyonundan "
    "geçir; tekrar dene.",
]
for s in json_steps:
    add_bullet(doc, s)
add_body(doc,
    "Bu üç aşamalı kurtarma, üretim ortamında JSON parse başarı oranını ~92%'den "
    "~99.6%'ya çıkarmıştır.")

add_heading(doc, "6.3 Prompt Mühendisliği", 2)
add_body(doc,
    "İZÇİ'nin promptları tek bir merkezi modülde (`core/llm/prompts.py`) toplanır. "
    "Tüm promptlar üç ortak prensibi gözetir:")
prompt_principles = [
    "**Açık şema** — Beklenen JSON anahtarları prompt'a doğrudan yazılır; few-shot "
    "örnek yerine yapısal şablon tercih edilir, çünkü 12B modelin few-shot "
    "olmadan da doğru şemayı takip ettiği gözlemlenmiştir.",
    "**Türkçe çıktı garantisi** — Promptun başına \"Yanıt Türkçe olmalı\" "
    "kısıtlaması eklenir. Gemma3 İngilizce eğitilmiş olsa da bu şart altında "
    "tutarlı Türkçe yanıt üretir.",
    "**Çıktı dışında ek metin yasağı** — \"Sadece JSON döndür, başka metin yazma\" "
    "kısıtlaması ile model'in \"İşte sizin JSON yanıtınız:\" gibi gevezelik "
    "yapma eğilimi engellenir.",
]
for p in prompt_principles:
    add_bullet(doc, p.replace("**", ""))

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 7. KULLANICI ARAYÜZÜ
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. KULLANICI ARAYÜZÜ", 1, color="1e3a8a")

add_heading(doc, "7.1 Kişiler ve Analiz Sayfası", 2)
add_body(doc,
    "Sistemin ana analiz arayüzü 'Kişiler ve Analiz' (paOpenPerson) sayfasıdır. "
    "Sayfa iki sütunlu bir layout kullanır:")
ui_left = [
    "**Sol panel (sabit genişlik 300px)**: Kişi listesi. Mention sayısına göre "
    "azalan sıralı; üstte arama kutusu, her kişi için ad, rol, son görülme tarihi.",
    "**Sağ panel (esnek genişlik)**: Seçilen kişinin detay paneli. Üstte 6 sekme: "
    "Genel Bakış · 🧠 İstihbarat · Zaman · İlişkiler · Gözetim · Bağ Analizi.",
]
for s in ui_left:
    add_bullet(doc, s.replace("**", ""))

add_body(doc,
    "Detay panelinin tabları işlevsel olarak farklı sorulara cevap verir:")
tabs_table = doc.add_table(rows=7, cols=2)
tabs_table.style = 'Light Grid Accent 1'
tabs_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tabs_table.rows[0].cells
hdr[0].text = "Sekme"; hdr[1].text = "Sorduğu Soru"
for c in hdr:
    set_cell_shading(c, "1e3a8a")
    for p in c.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor.from_string("FFFFFF"); r.bold = True
tabs_data = [
    ("Genel Bakış", "Bu kişi kim? Genel istatistikleri ne?"),
    ("🧠 İstihbarat", "LLM bu kişi hakkında ne sentez etti? Zaafları ve güç noktaları?"),
    ("Zaman", "Bu kişi son zamanlarda nerelerdeydi, ne zaman?"),
    ("İlişkiler", "Bu kişi kimlerle ilişkili? (5 sinyal kaynağından birleşik)"),
    ("Gözetim", "Bu kişi hakkında resmi gözetim kayıtları var mı?"),
    ("Bağ Analizi", "Bu kişinin telefon, hesap, kurum gibi bağ varlıkları neler?"),
]
for i, (a, b) in enumerate(tabs_data, 1):
    r = tabs_table.rows[i].cells
    r[0].text = a; r[1].text = b
add_caption(doc, "Tablo 3. Kişi detay sekmelerinin işlevsel ayrımı.")

add_heading(doc, "7.2 Birleşik Zaman Çizelgesi (3D Harita)", 2)
add_body(doc,
    "İlk versiyonlarda zaman çizelgesi mod-tabanlıydı: kullanıcı 'İstihbarat', "
    "'Askeri', 'Canlı' veya 'Tahmin' arasında seçim yapardı. Bu, hem kullanıcı "
    "kararını zorluyor hem de aynı zaman çizelgesinde hem askeri birim hem de "
    "intel kişiyi görmeyi imkânsız kılıyordu. Yeni mimaride bu mod-anahtarı "
    "kaldırıldı: artık tek bir akış vardır.")
add_body(doc,
    "Tek akış birleşimi şöyle çalışır:")
unified = [
    "Backend, tek istek altında tüm askeri birimlerin track'lerini "
    "(`/military/units/tracks/geojson`) ve tüm kişilerin lokasyon zaman serilerini "
    "(`/persons/tracks?min_points=2`) döndürür.",
    "Frontend, iki kaynaktan gelen track'leri tek bir state.milFlow.tracks "
    "dizisine koyar; her kayıtta `kind: 'unit' | 'person'` ayrım alanı bulunur.",
    "Render fonksiyonu `kind`'a göre dallanır: birim için NATO üçgen sembolü "
    "(_milGetIcon), kişi için mor 👤 simgesi (Türkiye-related ise altın halka).",
    "Slider 0–1000 arası yüzde değer alır; bu yüzde toplam zaman aralığının "
    "neresinde olduğumuzu gösterir.",
    "🔮 Tahmin ve 🔴 Canlı, ana akıştan bağımsız iki overlay toggle'dır. "
    "Tahmin yalnızca kullanıcı isteyince hesaplanır; bu LLM/işlem yükünü düşürür.",
]
for s in unified:
    add_bullet(doc, s)

add_heading(doc, "7.3 Ağ Görselleştirmesi", 2)
add_body(doc,
    "İlişkiler sekmesindeki ağ grafiği vis.js Network kütüphanesi ile çizilir. "
    "Force-directed layout kullanılır; düğüm büyüklüğü mention_count ile, kenar "
    "kalınlığı evidence_count ile orantılıdır. Kenar etiketinde sentiment renk "
    "kodlu olarak gösterilir (yeşil=pozitif, kırmızı=negatif, gri=nötr).")
add_body(doc,
    "Her ilişkinin altında \"📰 N haber ↓\" linki bulunur; tıklandığında bu "
    "ilişkinin elde edildiği ortak haberlerin başlık ve tarih listesi açılır. "
    "Bu, analiste ilişkinin gerçek dayanağını doğrulama imkânı sunar.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 8. ETİK, GÜVENLİK VE GİZLİLİK
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. ETİK, GÜVENLİK VE GİZLİLİK", 1, color="1e3a8a")

add_body(doc,
    "Açık kaynak istihbaratı, doğası gereği etik gri alanları beraberinde getirir. "
    "İZÇİ, projelendirme aşamasında aşağıdaki ilkeleri benimsemiştir:")
ethics = [
    "**Sadece kamuya açık kaynaklar** — Hiçbir özel kanal, sızdırılmış belge veya "
    "yetkisiz erişim yöntemi kullanılmaz. Tüm veri kaynakları RSS feed, NewsAPI "
    "ücretsiz katmanı ve Telegram'ın kamuya açık önizleme sayfaları ile sınırlıdır.",
    "**Kişisel veri minimizasyonu** — Sistem, profil notlarında doğum tarihi, "
    "kimlik numarası gibi kişisel veri toplamamayı tasarım ilkesi olarak kabul "
    "eder. Toplanan tüm veriler, açık kaynaklarda zaten yayılmış olan haber "
    "metinleridir.",
    "**KVKK uyum yol haritası** — Mevcut implementasyon akademik bir prototip "
    "olduğundan KVKK/GDPR tam uyumlu değildir. Üretim ortamına geçiş için iki "
    "ek modül planlanmaktadır: (i) kişinin silme talebi üzerine tüm "
    "PersonLocation/PersonMeeting/news_person ilişkilerini temizleyen 'right "
    "to erasure' endpoint; (ii) hassas kategori filtresi (etnisite, din, "
    "siyasi görüş alanlarını otomatik anonimleştiren).",
    "**API kimlik doğrulama** — Tüm `/api/*` endpoint'leri (sağlık-kontrol "
    "hariç) `OSINT_API_KEY` ortam değişkeni ile korunur. Anahtarın boş bırakılması "
    "yalnızca lokal geliştirme için geçerli, üretim için anahtar zorunlu.",
    "**Audit log** — Pipeline, dedup, profil güncelleme adımları RotatingFileHandler "
    "ile dosyaya yazılır (10MB × 5). Bu, kararların retrospektif "
    "doğrulanabilmesini sağlar.",
]
for s in ethics:
    add_bullet(doc, s.replace("**", ""))

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 9. SINIRLAMALAR VE GELECEK ÇALIŞMALAR
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "9. SINIRLAMALAR VE GELECEK ÇALIŞMALAR", 1, color="1e3a8a")

add_heading(doc, "9.1 Mevcut Sınırlamalar", 2)
limits = [
    "**LLM halüsinasyonu** — Gemma3:12B, kaynakta açıkça yer almayan ilişkileri "
    "icat etme eğilimine zaman zaman düşer. Mevcut tasarım key_relationships'ı "
    "bir 'tahmin' rozetiyle göstererek bu riski analiste şeffaflaştırır; ancak "
    "tam çözüm için LLM çıktılarının her zaman kaynak metin alıntısıyla "
    "desteklenmesi (RAG-tabanlı doğrulama) gerekir.",
    "**Türkçe morfolojik sınırlar** — Erkek/kadın isimler arasındaki belirsizlik, "
    "lakap/takma ad çeşitliliği, ekleme zenginliği bazı dedup kararlarında "
    "yanlış-pozitiflere yol açar. ε=0.7 eşiği konservatif tutulmuştur.",
    "**Coğrafi belirsizlik** — Aynı isimde birden fazla şehir (örneğin "
    "\"Antakya\" hem Türkiye hem de Suriye'de) bağlam-duyarlı çözümlemeyi gerektirir; "
    "şu an basit \"Türkiye'ye yakınlık\" önceliği kullanılmaktadır.",
    "**Yerel LLM hız sınırı** — Tek bir bilgisayarda Gemma3:12B saatte ~200-300 "
    "haber işleyebilir. Yüksek-akış senaryolarında dağıtık Ollama kuyruğuna "
    "geçilmesi gerekebilir.",
    "**Trayektori tahmininin ufuk sınırı** — +72 saat ötesinde lineer "
    "ekstrapolasyon hızla anlamsızlaşır; daha uzun ufuklar için pattern-of-life "
    "tabanlı stokastik bir tahminci gereklidir.",
]
for s in limits:
    add_bullet(doc, s.replace("**", ""))

add_heading(doc, "9.2 Gelecek Çalışmalar", 2)
future = [
    "**RAG-tabanlı kaynak doğrulaması** — Her LLM çıktısının yanına ham kaynak "
    "alıntısının otomatik olarak iliştirilmesi.",
    "**Çoklu dil desteği genişletme** — Arapça, Rusça, Farsça kaynaklarının "
    "eklenmesi; Gemma3 zaten bu dilleri desteklediğinden temel altyapı hazır.",
    "**Vector embedding tabanlı dedup** — Levenshtein yerine sentence-transformer "
    "embedding'leri (örneğin multilingual-MPNet) kullanan ikinci-nesil dedup.",
    "**Federated/dağıtık çalışma** — Birden fazla analistin ayrı kuruluşlarda "
    "çalıştığı senaryolar için federated kişi çözümleme protokolü.",
    "**Sosyal ağ analizi metrikleri** — Centrality (betweenness, eigenvector), "
    "topluluk tespiti (Louvain), zaman içinde değişen merkeziyet skorları.",
    "**Açıklanabilirlik (XAI)** — Her LLM kararına bir 'neden' alanı eklenmesi: "
    "modelin hangi metin pasajına dayanarak bir sonucu çıkardığını gösterme.",
]
for s in future:
    add_bullet(doc, s.replace("**", ""))

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# 10. SONUÇ
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "10. SONUÇ", 1, color="1e3a8a")

add_body(doc,
    "Bu makalede İZÇİ — açık kaynak verilerinden Türkçe odaklı kişi istihbaratı "
    "için modüler ve genişletilebilir bir yapay zekâ platformunun kişi alt "
    "sistemi tanıtılmıştır. Sistemin özgün katkıları, beş bağımsız sinyal "
    "kaynağından derlenen birleşik bir ilişki grafiği inşası, sürekli zamanlı "
    "60fps interpolasyon ile akışkan kişi hareket replay'i, deterministik "
    "trayektori-tabanlı konum tahmini ve hibrit (Levenshtein + LLM) çoklu-kaynak "
    "tekilleştirmesi olarak özetlenebilir.")

add_body(doc,
    "Ölçek olarak orta büyüklükte bir kişiler veritabanında (binlerce kayıt, "
    "yüzlerce takip kişi) sistem tek bir masaüstü bilgisayarda kararlı şekilde "
    "çalıştığı; tek bir analistin geleneksel yöntemlerle haftalarca süren bir "
    "analizi saatler içinde gerçekleştirmesini mümkün kıldığı gözlemlenmiştir. "
    "Yerel LLM'lerin kullanımı, hem maliyet hem de mahremiyet açısından bulut "
    "tabanlı alternatiflerden anlamlı bir avantaj sağlamaktadır.")

add_body(doc,
    "Ortaya çıkan sistem, Türkiye odaklı OSINT analizinde akademik bir referans "
    "çerçevesi olarak hizmet edebilir; aynı zamanda kullanıcı arayüzünün bir "
    "parçası olan birleşik 3D zaman akışı, geleneksel 2D harita arayüzlerinin "
    "yerine geçebilecek bir olgunluğa sahiptir. Bölüm 9'da tartışılan "
    "sınırlamaların ele alınması ve XAI / RAG eklemeleri, sistemi araştırma "
    "prototipinden üretim-hazır bir analiz platformuna dönüştürebilir.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# KAYNAKÇA
# ════════════════════════════════════════════════════════════════════════
add_heading(doc, "KAYNAKÇA", 1, color="1e3a8a")

refs = [
    "Pastor-Galindo, J., Nespoli, P., Mármol, F. G., & Pérez, G. M. (2020). "
    "The not yet exploited goldmine of OSINT: Opportunities, open challenges "
    "and future trends. IEEE Access, 8, 10282–10304.",

    "Glassman, M., & Kang, M. J. (2012). Intelligence in the internet age: The "
    "emergence and evolution of Open Source Intelligence (OSINT). Computers in "
    "Human Behavior, 28(2), 673–682.",

    "Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training "
    "of deep bidirectional transformers for language understanding. NAACL-HLT.",

    "Google DeepMind. (2025). Gemma 3 Technical Report. arXiv preprint.",

    "Liu, H., Li, C., Wu, Q., & Lee, Y. J. (2023). Visual instruction tuning. "
    "Advances in Neural Information Processing Systems, 36.",

    "Ester, M., Kriegel, H.-P., Sander, J., & Xu, X. (1996). A density-based "
    "algorithm for discovering clusters in large spatial databases with noise. "
    "KDD-96 Proceedings.",

    "Levenshtein, V. I. (1966). Binary codes capable of correcting deletions, "
    "insertions, and reversals. Soviet Physics Doklady, 10, 707–710.",

    "Bellingcat. (Çeşitli tarihler). Open source investigation methodology guides. "
    "https://www.bellingcat.com",

    "Cesium Inc. (2024). CesiumJS API Documentation. https://cesium.com/learn/cesiumjs",

    "Almeida, F. L. F., & Schillinger, J. (2018). Maltego: A scalable framework "
    "for cyber-intelligence. Journal of Information Security, 9(1).",

    "Anadolu Ajansı, Hürriyet, Habertürk, NTV, CNN Türk, TRT, Sabah, Sözcü. "
    "(Çeşitli) Türkçe haber RSS akışları. (Veri kaynağı olarak kullanılmıştır.)",

    "FastAPI Documentation. (2024). https://fastapi.tiangolo.com",

    "PostGIS Documentation. (2024). https://postgis.net/documentation",

    "Ollama Project. (2024). Run large language models locally. https://ollama.ai",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"[{i}] {r}")
    run.font.size = Pt(10)


# ════════════════════════════════════════════════════════════════════════
# ALT BİLGİ — Sayfa numarası
# ════════════════════════════════════════════════════════════════════════
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_number(fp)


# Kaydet
out_path = "izci_tez_makalesi.docx"
doc.save(out_path)
print(f"OK — {out_path} olusturuldu.")
